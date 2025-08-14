"""
Enhanced export and save functionality for different file formats
"""
import json
import os
from datetime import datetime
from pathlib import Path
import pandas as pd
from typing import Dict, Any, List
import tempfile

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
except ImportError:
    SimpleDocTemplate = None

import structlog

logger = structlog.get_logger()

def create_export_directory():
    """Create export directory if it doesn't exist."""
    export_dir = Path("./exports")
    export_dir.mkdir(exist_ok=True)
    return export_dir

def export_to_docx(doc_info: Dict, analysis_result: Dict, chat_history: List) -> str:
    """Export analysis to DOCX format."""
    try:
        if Document is None:
            return "خطأ: مكتبة python-docx غير مثبتة"
        
        export_dir = create_export_directory()
        filename = f"تحليل_{doc_info['doc_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = export_dir / filename
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(f"تقرير تحليل المستند: {doc_info['filename']}", 0)
        title.alignment = 2  # Right alignment for Arabic
        
        # Document info
        doc.add_heading('معلومات المستند', level=1)
        doc.add_paragraph(f"اسم الملف: {doc_info['filename']}")
        doc.add_paragraph(f"معرف المستند: {doc_info['doc_id']}")
        doc.add_paragraph(f"حجم النص: {len(doc_info['text']):,} حرف")
        doc.add_paragraph(f"عدد الكلمات: {len(doc_info['text'].split()):,}")
        doc.add_paragraph(f"وقت التحليل: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Analysis results
        if analysis_result and analysis_result.get('success'):
            doc.add_heading('نتائج التحليل', level=1)
            
            if analysis_result.get('analysis'):
                analysis = analysis_result['analysis']
                doc.add_heading('النتائج الرئيسية', level=2)
                for i, finding in enumerate(analysis.findings, 1):
                    doc.add_paragraph(f"{i}. {finding}")
                
                doc.add_paragraph(f"التبرير: {analysis.rationale}")
                doc.add_paragraph(f"يحتاج تحليل إحصائي: {'نعم' if analysis.needs_scipy else 'لا'}")
            
            if analysis_result.get('decision'):
                decision = analysis_result['decision']
                doc.add_heading('القرار النهائي', level=2)
                doc.add_paragraph(f"التصنيف: {decision.label}")
                doc.add_paragraph(f"مستوى الثقة: {decision.confidence:.0%}")
                doc.add_paragraph("المعايير المستخدمة:")
                for criterion in decision.criteria:
                    doc.add_paragraph(f"• {criterion}")
            
            if analysis_result.get('scipy_out') and analysis_result['scipy_out'].get('analysis_performed'):
                doc.add_heading('التحليل الإحصائي', level=2)
                scipy_results = analysis_result['scipy_out']
                doc.add_paragraph(f"التحليلات المنجزة: {', '.join(scipy_results['analysis_performed'])}")
        
        # Chat history
        if chat_history:
            doc.add_heading('سجل المحادثة', level=1)
            for i, message in enumerate(chat_history, 1):
                if message['role'] == 'user':
                    doc.add_paragraph(f"السؤال {i}: {message['content']}")
                else:
                    doc.add_paragraph(f"الإجابة {i}: {message['content']}")
                doc.add_paragraph("")  # Add spacing
        
        # Save document
        doc.save(str(filepath))
        logger.info("DOCX export completed", filepath=str(filepath))
        return str(filepath)
        
    except Exception as e:
        logger.error("DOCX export failed", error=str(e))
        return f"خطأ في التصدير: {str(e)}"

def export_to_pdf_report(doc_info: Dict, analysis_result: Dict, chat_history: List) -> str:
    """Export analysis to PDF format using ReportLab."""
    try:
        if SimpleDocTemplate is None:
            return "خطأ: مكتبة reportlab غير مثبتة"
        
        export_dir = create_export_directory()
        filename = f"تقرير_{doc_info['doc_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        filepath = export_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        styles = getSampleStyleSheet()
        
        # Arabic-friendly style
        arabic_style = ParagraphStyle(
            'Arabic',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=12,
            alignment=2,  # Right alignment
            spaceAfter=12
        )
        
        story = []
        
        # Title
        title = Paragraph(f"تقرير تحليل المستند: {doc_info['filename']}", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Document info
        story.append(Paragraph("معلومات المستند", styles['Heading1']))
        story.append(Paragraph(f"اسم الملف: {doc_info['filename']}", arabic_style))
        story.append(Paragraph(f"معرف المستند: {doc_info['doc_id']}", arabic_style))
        story.append(Paragraph(f"حجم النص: {len(doc_info['text']):,} حرف", arabic_style))
        story.append(Paragraph(f"عدد الكلمات: {len(doc_info['text'].split()):,}", arabic_style))
        story.append(Spacer(1, 12))
        
        # Analysis results
        if analysis_result and analysis_result.get('success'):
            story.append(Paragraph("نتائج التحليل", styles['Heading1']))
            
            if analysis_result.get('analysis'):
                analysis = analysis_result['analysis']
                story.append(Paragraph("النتائج الرئيسية", styles['Heading2']))
                for i, finding in enumerate(analysis.findings, 1):
                    story.append(Paragraph(f"{i}. {finding}", arabic_style))
                
                story.append(Paragraph(f"التبرير: {analysis.rationale}", arabic_style))
                story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        logger.info("PDF export completed", filepath=str(filepath))
        return str(filepath)
        
    except Exception as e:
        logger.error("PDF export failed", error=str(e))
        return f"خطأ في التصدير: {str(e)}"

def export_to_excel(doc_info: Dict, analysis_result: Dict, chat_history: List) -> str:
    """Export analysis to Excel format."""
    try:
        export_dir = create_export_directory()
        filename = f"بيانات_{doc_info['doc_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        filepath = export_dir / filename
        
        with pd.ExcelWriter(str(filepath), engine='openpyxl') as writer:
            # Document info sheet
            doc_data = {
                'المعلومة': ['اسم الملف', 'معرف المستند', 'حجم النص', 'عدد الكلمات', 'وقت التحليل'],
                'القيمة': [
                    doc_info['filename'],
                    doc_info['doc_id'], 
                    f"{len(doc_info['text']):,} حرف",
                    f"{len(doc_info['text'].split()):,}",
                    datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                ]
            }
            pd.DataFrame(doc_data).to_excel(writer, sheet_name='معلومات المستند', index=False)
            
            # Analysis results sheet
            if analysis_result and analysis_result.get('success'):
                analysis_data = []
                if analysis_result.get('analysis'):
                    analysis = analysis_result['analysis']
                    for i, finding in enumerate(analysis.findings, 1):
                        analysis_data.append({'الرقم': i, 'النتيجة': finding})
                    
                if analysis_data:
                    pd.DataFrame(analysis_data).to_excel(writer, sheet_name='نتائج التحليل', index=False)
                
                # Decision sheet
                if analysis_result.get('decision'):
                    decision = analysis_result['decision']
                    decision_data = {
                        'المعيار': ['التصنيف', 'مستوى الثقة'] + [f'معيار {i+1}' for i in range(len(decision.criteria))],
                        'القيمة': [decision.label, f"{decision.confidence:.0%}"] + decision.criteria
                    }
                    pd.DataFrame(decision_data).to_excel(writer, sheet_name='القرار النهائي', index=False)
            
            # Chat history sheet
            if chat_history:
                chat_data = []
                for i, message in enumerate(chat_history, 1):
                    chat_data.append({
                        'الرقم': i,
                        'النوع': 'سؤال' if message['role'] == 'user' else 'إجابة',
                        'المحتوى': message['content']
                    })
                pd.DataFrame(chat_data).to_excel(writer, sheet_name='سجل المحادثة', index=False)
        
        logger.info("Excel export completed", filepath=str(filepath))
        return str(filepath)
        
    except Exception as e:
        logger.error("Excel export failed", error=str(e))
        return f"خطأ في التصدير: {str(e)}"

def export_chat_to_text(chat_history: List, doc_info: Dict) -> str:
    """Export chat history to simple text file."""
    try:
        export_dir = create_export_directory()
        filename = f"محادثة_{doc_info['doc_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = export_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"سجل محادثة - المستند: {doc_info['filename']}\n")
            f.write(f"التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 50 + "\n\n")
            
            for i, message in enumerate(chat_history, 1):
                if message['role'] == 'user':
                    f.write(f"السؤال {i}: {message['content']}\n\n")
                else:
                    f.write(f"الإجابة {i}: {message['content']}\n\n")
                f.write("-" * 30 + "\n\n")
        
        logger.info("Text export completed", filepath=str(filepath))
        return str(filepath)
        
    except Exception as e:
        logger.error("Text export failed", error=str(e))
        return f"خطأ في التصدير: {str(e)}"

def get_available_export_formats():
    """Get list of available export formats."""
    formats = [
        ("JSON", "json", "متاح دائماً"),
        ("Excel", "xlsx", "متاح دائماً"),
        ("نص عادي", "txt", "متاح دائماً")
    ]
    
    if Document is not None:
        formats.append(("Word Document", "docx", "متاح"))
    else:
        formats.append(("Word Document", "docx", "غير متاح - pip install python-docx"))
    
    if SimpleDocTemplate is not None:
        formats.append(("PDF Report", "pdf", "متاح"))
    else:
        formats.append(("PDF Report", "pdf", "غير متاح - pip install reportlab"))
    
    return formats
